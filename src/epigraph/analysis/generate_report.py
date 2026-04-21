"""Generate a comprehensive Word document report from all analysis results.

Uses python-docx for reliable Word compatibility. Reads all processed
analysis outputs and produces a consolidated .docx report.

Usage::

    python -m epigraph.analysis.generate_report
    python -m epigraph.analysis.generate_report --output reports/my_report.docx
"""

from __future__ import annotations

from pathlib import Path
from typing import TYPE_CHECKING, Any

import click
import polars as pl
from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor

from epigraph.common.logging import get_logger

if TYPE_CHECKING:
    # ``docx.Document`` is a factory function; the concrete class lives at
    # ``docx.document.Document``. Use the class for static type annotations
    # so attribute access on ``doc.styles`` / ``doc.add_table`` is checkable.
    from docx.document import Document as DocxDocument

log = get_logger(__name__)

# ---------------------------------------------------------------------------
# Styling helpers
# ---------------------------------------------------------------------------

BLUE = RGBColor(0x1B, 0x4F, 0x72)
GREY = RGBColor(0x7F, 0x8C, 0x8D)


def _style_heading(doc: DocxDocument) -> None:
    """Configure heading styles."""
    for i, (size, colour) in enumerate([(18, BLUE), (14, RGBColor(0x2C, 0x3E, 0x50))], 1):
        style = doc.styles[f"Heading {i}"]
        style.font.size = Pt(size)
        style.font.color.rgb = colour
        style.font.name = "Arial"
        style.font.bold = True


def _add_table(doc: DocxDocument, headers: list[str], rows: list[list[str]]) -> None:
    """Add a formatted table to the document."""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Light Grid Accent 1"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Header row
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            for run in p.runs:
                run.bold = True
                run.font.size = Pt(9)

    # Data rows
    for row_idx, row_data in enumerate(rows):
        for col_idx, val in enumerate(row_data):
            cell = table.rows[row_idx + 1].cells[col_idx]
            cell.text = str(val)
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.size = Pt(9)


# ---------------------------------------------------------------------------
# Data loading (reused from before)
# ---------------------------------------------------------------------------


def _load_comparison(path: Path) -> dict[str, Any]:
    df = pl.read_parquet(path)
    n_sig = df.filter(pl.col("significant")).height if "significant" in df.columns else 0
    n_nominal = df.filter(pl.col("p_value") < 0.05).height
    top = df.filter(pl.col("significant")).sort("p_value").head(15)
    return {
        "label": path.stem,
        "n_genes": df.height,
        "n_nominal": n_nominal,
        "n_significant": n_sig,
        "top_hits": [
            {
                "gene": row["feature"],
                "cohens_d": f"{row['cohens_d']:+.3f}" if row["cohens_d"] == row["cohens_d"] else "N/A",
                "delta_mean": f"{row['delta_mean']:.4f}" if row["delta_mean"] == row["delta_mean"] else "N/A",
                "p_value": f"{row['p_value']:.2e}",
                "q_value": f"{row['q_value']:.6f}",
            }
            for row in top.iter_rows(named=True)
        ],
    }


def _load_enrichment(path: Path, names_path: Path | None = None) -> dict[str, Any]:
    df = pl.read_parquet(path)
    n_sig = df.filter(pl.col("significant")).height if "significant" in df.columns else 0
    sig_df = df.filter(pl.col("significant")).sort("p_value") if n_sig > 0 else df.head(0)

    if names_path and names_path.exists():
        names = pl.read_parquet(names_path)
        id_col = "pathway_id" if "pathway_id" in names.columns else names.columns[0]
        name_col = "pathway_name" if "pathway_name" in names.columns else names.columns[1]
        sig_df = sig_df.join(names, left_on="pathway", right_on=id_col, how="left")
    else:
        name_col = "pathway"

    return {
        "source": path.stem,
        "n_tested": df.height,
        "n_significant": n_sig,
        "pathways": [
            {
                "name": str(row.get(name_col, row.get("pathway", "")))[:60],
                "odds_ratio": f"{row.get('odds_ratio', 0):.2f}" if row.get("odds_ratio") else "N/A",
                "q_value": f"{row['q_value']:.4f}",
                "n_overlap": str(row.get("n_overlap", row.get("n_hits", 0))),
            }
            for row in sig_df.head(15).iter_rows(named=True)
        ],
    }


# ---------------------------------------------------------------------------
# Report builder
# ---------------------------------------------------------------------------


def build_report(
    output_path: Path,
    comparisons_dir: Path,
    enrichment_dir: Path,
    hms_dir: Path,
    external_dir: Path,
    mapping_path: Path,
    island_path: Path,
    figures_dir: Path,
    clinical_path: Path,
) -> None:
    """Build the Word document report."""
    doc = Document()
    _style_heading(doc)

    # Set default font
    style = doc.styles["Normal"]
    style.font.name = "Arial"
    style.font.size = Pt(11)

    # ---- Title page ----
    doc.add_paragraph()
    doc.add_paragraph()
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("Methylation Knowledge Graph")
    run.font.size = Pt(28)
    run.font.color.rgb = BLUE
    run.bold = True

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("CRC Biomarker Discovery Report")
    run.font.size = Pt(18)
    run.font.color.rgb = RGBColor(0x2C, 0x3E, 0x50)

    from datetime import date

    date_para = doc.add_paragraph()
    date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = date_para.add_run(f"Generated: {date.today().isoformat()}")
    run.font.size = Pt(10)
    run.font.color.rgb = GREY

    doc.add_page_break()

    # ---- 1. Executive Summary ----
    doc.add_heading("1. Executive Summary", level=1)

    # Load primary comparison for summary
    crc_path = comparisons_dir / "CRC_vs_Control.parquet"
    crc_data = _load_comparison(crc_path) if crc_path.exists() else None

    n_sig = crc_data["n_significant"] if crc_data else 0
    doc.add_paragraph(
        f"This report presents genome-wide DNA methylation analysis comparing colorectal cancer (CRC) "
        f"samples against healthy controls and polyps, "
        f"aggregated to {crc_data['n_genes']:,} gene-level features." if crc_data else
        "Comparison data not available."
    )
    doc.add_paragraph(
        f"Key findings: {n_sig:,} genes showed statistically significant differential methylation "
        f"(FDR < 0.05) between CRC and Control samples."
    )

    # ---- 2. Cohort Description ----
    doc.add_page_break()
    doc.add_heading("2. Cohort Description", level=1)

    if clinical_path.exists():
        clin = pl.read_parquet(clinical_path)
        doc.add_paragraph(
            f"The study cohort comprises {len(clin):,} samples with clinical metadata "
            f"derived from CLIN_* worksheets in the biomarker aggregation spreadsheet."
        )
        cats = clin.group_by("clinical_category").len().sort("len", descending=True)
        _add_table(doc,
            ["Clinical Category", "Sample Count"],
            [[row["clinical_category"], f"{row['len']:,}"] for row in cats.iter_rows(named=True)],
        )

    # ---- 3. Genomic Annotation ----
    doc.add_page_break()
    doc.add_heading("3. Genomic Annotation", level=1)

    if mapping_path.exists():
        doc.add_heading("3.1 CpG-to-Gene Mapping", level=2)
        mapping = pl.read_parquet(mapping_path)
        n_cpgs = mapping["cpg_id"].n_unique()
        n_genes = mapping["gene_symbol"].n_unique()
        doc.add_paragraph(
            f"CpG sites were mapped to GENCODE v45 genes using coordinate overlap. "
            f"{n_cpgs:,} unique CpGs mapped to {n_genes:,} genes via "
            f"{len(mapping):,} total mappings."
        )
        overlap = mapping.group_by("overlap_type").len().sort("len", descending=True)
        _add_table(doc,
            ["Overlap Type", "Count"],
            [[row["overlap_type"], f"{row['len']:,}"] for row in overlap.iter_rows(named=True)],
        )

    if island_path.exists():
        doc.add_heading("3.2 CpG Island Context", level=2)
        ctx = pl.read_parquet(island_path)
        total = ctx.height
        ctx_counts = ctx.group_by("context").len().sort("len", descending=True)
        _add_table(doc,
            ["Context", "Count", "Percentage"],
            [[row["context"], f"{row['len']:,}", f"{row['len']/total*100:.1f}%"]
             for row in ctx_counts.iter_rows(named=True)],
        )

    # ---- 4. Differential Methylation ----
    doc.add_page_break()
    doc.add_heading("4. Differential Methylation", level=1)
    doc.add_paragraph(
        "Gene-level methylation was computed as the mean beta value across all CpGs mapping "
        "to each gene. Cohort comparisons used the Mann-Whitney U test with Benjamini-Hochberg "
        "FDR correction at alpha = 0.05."
    )

    for i, comp_file in enumerate(sorted(comparisons_dir.glob("*.parquet")), 1):
        comp = _load_comparison(comp_file)
        doc.add_heading(f"4.{i} {comp['label'].replace('_', ' ')}", level=2)
        doc.add_paragraph(
            f"{comp['n_genes']:,} genes tested. "
            f"{comp['n_nominal']:,} nominally significant (p < 0.05). "
            f"{comp['n_significant']:,} FDR-significant (q < 0.05)."
        )
        if comp["top_hits"]:
            _add_table(doc,
                ["Gene", "Cohen\u2019s d", "Delta Mean", "p-value", "q-value"],
                [[h["gene"], h["cohens_d"], h["delta_mean"], h["p_value"], h["q_value"]]
                 for h in comp["top_hits"]],
            )
            doc.add_paragraph()

    # ---- 5. Pathway Enrichment ----
    doc.add_page_break()
    doc.add_heading("5. Pathway Enrichment", level=1)
    doc.add_paragraph(
        "Differentially methylated genes were tested for enrichment in Reactome pathways "
        "and GO terms using Fisher\u2019s exact test with FDR correction."
    )

    pw_names = external_dir / "reactome_pathways.parquet"
    for i, enrich_file in enumerate(sorted(enrichment_dir.glob("*.parquet")), 1):
        names = pw_names if "reactome" in enrich_file.name else None
        enrich = _load_enrichment(enrich_file, names)

        doc.add_heading(f"5.{i} {enrich['source'].replace('_', ' ')}", level=2)
        doc.add_paragraph(
            f"{enrich['n_tested']:,} pathways/terms tested. "
            f"{enrich['n_significant']} FDR-significant."
        )
        if enrich["pathways"]:
            _add_table(doc,
                ["Pathway/Term", "Odds Ratio", "q-value", "Gene Overlap"],
                [[p["name"], p["odds_ratio"], p["q_value"], p["n_overlap"]]
                 for p in enrich["pathways"]],
            )
            doc.add_paragraph()

    # ---- 6. Hypermethylation Scoring ----
    doc.add_page_break()
    doc.add_heading("6. Hypermethylation Scoring", level=1)

    hms_loaded = False
    for q in ["0.99", "0.95", "0.999"]:
        hms_path = hms_dir / f"hms_scores_q{q.replace('.', '_')}.parquet"
        if hms_path.exists():
            df = pl.read_parquet(hms_path)
            doc.add_paragraph(
                f"Hypermethylation was scored using control-quantile thresholds (q={q}). "
                f"For each gene, a per-gene threshold was set at the {float(q)*100:.0f}th "
                f"percentile of control samples. The HMS count is the number of genes "
                f"exceeding their threshold per sample."
            )
            if "clinical_category" in df.columns:
                stats = (
                    df.filter(pl.col("clinical_category").is_in(["CRC", "Control", "polyps"]))
                    .group_by("clinical_category")
                    .agg(
                        pl.col("hms_count").mean().alias("mean"),
                        pl.col("hms_count").median().alias("median"),
                        pl.len().alias("n"),
                    )
                    .sort("clinical_category")
                )
                _add_table(doc,
                    ["Category", "N", "Mean HMS", "Median HMS"],
                    [[row["clinical_category"], str(row["n"]),
                      f"{row['mean']:.0f}", f"{row['median']:.0f}"]
                     for row in stats.iter_rows(named=True)],
                )
            hms_loaded = True
            break

    if not hms_loaded:
        doc.add_paragraph("Hypermethylation scoring data not available.")

    # ---- 7. Figures ----
    doc.add_page_break()
    doc.add_heading("7. Figures", level=1)

    fig_descriptions = {
        "volcano_CRC_vs_Control": "Volcano plot: gene-level differential methylation, CRC vs Control. Red = FDR-significant.",
        "volcano_CRC_vs_polyps": "Volcano plot: gene-level differential methylation, CRC vs polyps.",
        "volcano_polyps_vs_Control": "Volcano plot: gene-level differential methylation, polyps vs Control.",
        "dotplot_reactome": "Reactome pathway enrichment dot plot. Size = gene overlap, colour = odds ratio.",
        "hms_distribution": "Hypermethylation score distribution by clinical category.",
        "heatmap": "Top 50 differentially methylated genes (z-scored), grouped by clinical category.",
    }

    for fig_path in sorted(figures_dir.resolve().glob("*.png")):
        fig_name = fig_path.stem
        doc.add_heading(fig_name.replace("_", " "), level=2)

        # Find matching description
        for key, desc in fig_descriptions.items():
            if key in fig_name:
                doc.add_paragraph(desc)
                break

        doc.add_picture(str(fig_path), width=Inches(5.5))
        last_para = doc.paragraphs[-1]
        last_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.add_paragraph()

    # ---- 8. Methods ----
    doc.add_page_break()
    doc.add_heading("8. Methods", level=1)

    doc.add_heading("8.1 Data Processing", level=2)
    doc.add_paragraph(
        "The beta-matrix CSV is converted to per-chromosome Parquet files using single-pass "
        "line-by-line extraction with pre-allocated numpy float32 arrays, keeping peak memory "
        "low regardless of matrix size."
    )

    doc.add_heading("8.2 Gene-Level Aggregation", level=2)
    doc.add_paragraph(
        "CpG beta values were aggregated to gene level by computing the arithmetic mean of "
        "all CpGs mapping to each gene (promoter and gene body regions). A CpG may contribute "
        "to multiple genes if gene loci overlap. Genes with fewer than 3 mapped CpGs were excluded."
    )

    doc.add_heading("8.3 Statistical Testing", level=2)
    doc.add_paragraph(
        "Group comparisons used the Mann-Whitney U test (two-sided) with Benjamini-Hochberg "
        "FDR correction at alpha = 0.05. Effect sizes reported as Cohen\u2019s d. "
        "Pathway enrichment used Fisher\u2019s exact test (overrepresentation analysis)."
    )

    doc.add_heading("8.4 Annotations", level=2)
    doc.add_paragraph(
        "Gene annotations: GENCODE v45 (GRCh38). Pathways: Reactome. Functional terms: GO/GOA. "
        "CpG islands: UCSC hg38 cpgIslandExt. All coordinates are 1-based."
    )

    # Save
    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_path))
    size_kb = output_path.stat().st_size / 1024
    log.info("report_saved", path=str(output_path), size_kb=round(size_kb))
    click.echo(f"Report written to {output_path} ({size_kb:.0f} KB)")


# ---------------------------------------------------------------------------
# CLI
# ---------------------------------------------------------------------------


@click.command("generate-report")
@click.option("--output", default="data/processed/reports/methylation_report.docx")
@click.option("--comparisons-dir", default="data/processed/comparisons_full")
@click.option("--enrichment-dir", default="data/processed/enrichment_full")
@click.option("--hms-dir", default="data/processed/hypermethylation")
@click.option("--external-dir", default="data/external")
@click.option("--mapping-path", default="data/processed/cpg_gene_mapping_full.parquet")
@click.option("--island-path", default="data/processed/cpg_island_context.parquet")
@click.option("--figures-dir", default="data/processed/figures")
@click.option("--clinical-path", default="data/processed/clinical_metadata.parquet")
def main(output: str, **kwargs: str) -> None:
    """Generate a comprehensive Word document report from all analysis results."""
    build_report(output_path=Path(output), **{k: Path(v) for k, v in kwargs.items()})


if __name__ == "__main__":
    main()
